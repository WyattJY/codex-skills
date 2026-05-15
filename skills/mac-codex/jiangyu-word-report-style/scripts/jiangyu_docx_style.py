# -*- coding: utf-8 -*-
"""Reusable python-docx helpers for Jiangyu's Word report style.

Copy this script into the active workspace, fill in `write_report`, and run it
with the bundled Codex Python runtime. It intentionally avoids project-specific
content so future reports can reuse the exact typography and table style.
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, east_asia="宋体", size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def set_paragraph_body(paragraph, first_line=True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(24 if first_line else 0)


def init_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    for name, east_asia, size in [("Heading 1", "黑体", 16), ("Heading 2", "黑体", 14)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size=16, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(4)
    fmt.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size=14, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_body(p, first_line=True)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_body_no_indent(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    set_paragraph_body(p, first_line=False)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def remove_all_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_border(cell, edge, val="single", sz="8", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    tag = "w:" + edge
    element = tc_borders.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        tc_borders.append(element)
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), sz)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold)


def apply_three_line_table(table):
    remove_all_borders(table)
    for cell in table.rows[0].cells:
        set_cell_border(cell, "top", sz="18")
        set_cell_border(cell, "bottom", sz="8")
    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", sz="18")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True)
    return p


def add_table(doc, caption, headers, rows, widths=None, font_size=10.5, center_cols=None):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    center_cols = set(center_cols or [])
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            set_cell_width(table.rows[0].cells[i], widths[i])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
            if widths:
                set_cell_width(cells[i], widths[i])
    apply_three_line_table(table)
    doc.add_paragraph()
    return table


def add_cover(doc, title, subtitle, info_rows):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(title)
    set_run_font(run, east_asia="黑体", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(subtitle)
    set_run_font(run, size=14, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    add_table(doc, "基本信息", ["项目", "内容"], info_rows, widths=[3.5, 12.0], center_cols=[0])
    doc.add_page_break()


def new_document():
    doc = Document()
    init_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.4)
    return doc


def write_report(doc):
    add_cover(
        doc,
        "报告标题",
        "报告副标题",
        [
            ["文档用途", "说明本文档用途"],
            ["适用项目", "项目名称"],
            ["生成日期", "YYYY-MM-DD"],
        ],
    )
    add_heading1(doc, "一、结论先行")
    add_body(doc, "在这里写正文。")
    add_heading1(doc, "二、技术说明")
    add_table(
        doc,
        "表1  示例三线表",
        ["字段", "说明"],
        [["字段一", "说明内容"], ["字段二", "说明内容"]],
        widths=[4.0, 11.5],
        center_cols=[0],
    )


def main():
    out = Path("jiangyu_report.docx")
    doc = new_document()
    write_report(doc)
    doc.save(out)
    print(out.resolve())


if __name__ == "__main__":
    main()
