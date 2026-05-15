import argparse
import re
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional


_DEFAULT_HEADING_REGEX = r"^#{1,6}\s*(参考文献|References)\s*$"
_CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff]")


@dataclass(frozen=True)
class _DecodedText:
    text: str
    encoding: str
    has_utf8_bom: bool
    newline: str


def _decode_text_file(path: Path, encoding: Optional[str]) -> _DecodedText:
    data = path.read_bytes()
    newline = "\r\n" if b"\r\n" in data else "\n"

    if encoding:
        text = data.decode(encoding)
        return _DecodedText(text=text, encoding=encoding, has_utf8_bom=False, newline=newline)

    if data.startswith(b"\xef\xbb\xbf"):
        return _DecodedText(
            text=data.decode("utf-8-sig"),
            encoding="utf-8-sig",
            has_utf8_bom=True,
            newline=newline,
        )

    for enc in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return _DecodedText(
                text=data.decode(enc),
                encoding=enc,
                has_utf8_bom=False,
                newline=newline,
            )
        except UnicodeDecodeError:
            continue

    # If we got here, something is very wrong; fall back to replacement to avoid crashing.
    return _DecodedText(
        text=data.decode("utf-8", errors="replace"),
        encoding="utf-8",
        has_utf8_bom=False,
        newline=newline,
    )


def _is_cjk_entry(entry: str) -> bool:
    return bool(_CJK_RE.search(entry))


def _strip_leading_marker(line: str) -> str:
    return re.sub(r"^\s*[-*]\s+", "", line).strip()


def _strip_leading_numbering(entry: str) -> str:
    # Examples: "1. ...", "12) ...", "[3] ...", "4、..."
    entry = re.sub(r"^\s*\[\s*\d+\s*\]\s*", "", entry)
    entry = re.sub(r"^\s*\d+\s*[.)、]\s*", "", entry)
    return entry.strip()


def _collapse_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _parse_entries(section_lines: Iterable[str]) -> list[str]:
    entries: list[str] = []
    buffer: list[str] = []

    for raw_line in section_lines:
        line = raw_line.strip()
        if not line:
            if buffer:
                entry = _collapse_spaces(" ".join(buffer))
                entry = _strip_leading_numbering(entry)
                if entry:
                    entries.append(entry)
                buffer = []
            continue

        buffer.append(_strip_leading_marker(raw_line))

    if buffer:
        entry = _collapse_spaces(" ".join(buffer))
        entry = _strip_leading_numbering(entry)
        if entry:
            entries.append(entry)

    return entries


def _sort_key_english(entry: str) -> str:
    # Most EN entries begin with author; normalize accents and case for stable A-Z ordering.
    normalized = unicodedata.normalize("NFKD", entry)
    normalized = "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")
    normalized = normalized.replace("’", "'").replace("“", '"').replace("”", '"')
    normalized = normalized.casefold()
    normalized = re.sub(r"^[^a-z]+", "", normalized)
    return normalized


def _find_reference_section(lines: list[str], heading_regex: str) -> tuple[int, int, int]:
    """
    Return (heading_line_index, section_start_index, section_end_index).
    section_start_index is the first line after the heading.
    section_end_index is the first line after the section (exclusive).
    """
    pattern = re.compile(heading_regex)
    heading_idx = -1
    heading_level = 1

    for i, line in enumerate(lines):
        if pattern.match(line):
            heading_idx = i
            heading_level = len(line) - len(line.lstrip("#"))

    if heading_idx < 0:
        raise ValueError(f"Could not find a reference section heading via regex: {heading_regex!r}")

    section_start = heading_idx + 1
    section_end = len(lines)
    next_heading = re.compile(rf"^#{{1,{heading_level}}}\s+")
    for j in range(section_start, len(lines)):
        if next_heading.match(lines[j]):
            section_end = j
            break

    return heading_idx, section_start, section_end


def _format_reference_section(entries: list[str]) -> list[str]:
    # One entry per paragraph, separated by a blank line.
    out: list[str] = [""]
    for entry in entries:
        out.append(entry)
        out.append("")
    return out


def _export_docx(
    *,
    entries: list[str],
    out_path: Path,
    title: str,
    font_en: str,
    font_cn: str,
    font_size_pt: int,
    hanging_cm: float,
) -> None:
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "python-docx is required for --docx-out. Install it with: python -m pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading(title, level=1)

    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(hanging_cm)
        p.paragraph_format.first_line_indent = Cm(-hanging_cm)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)

        run = p.add_run(entry)
        run.font.name = font_en
        run.font.size = Pt(font_size_pt)

        # Ensure Chinese characters use an East Asia font while Latin uses font_en.
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font_en)
        rfonts.set(qn("w:hAnsi"), font_en)
        rfonts.set(qn("w:eastAsia"), font_cn)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Format a thesis reference list in Markdown (CN first, EN second; EN A-Z by author) "
            "and optionally export it as a thesis-style Word (.docx) list."
        )
    )
    parser.add_argument("--md-in", required=True, help="Input Markdown file path.")
    parser.add_argument(
        "--heading-regex",
        default=_DEFAULT_HEADING_REGEX,
        help=f"Regex to match the reference section heading (default: {_DEFAULT_HEADING_REGEX!r}).",
    )
    parser.add_argument(
        "--encoding",
        default=None,
        help="Force input/output encoding (default: auto-detect; prefers UTF-8).",
    )

    write_group = parser.add_mutually_exclusive_group()
    write_group.add_argument(
        "--write-md",
        action="store_true",
        help="Write back to --md-in in-place.",
    )
    write_group.add_argument(
        "--md-out",
        default=None,
        help="Write formatted Markdown to this path (does not modify --md-in).",
    )

    parser.add_argument(
        "--docx-out",
        default=None,
        help="Optional output .docx path for the formatted reference list.",
    )
    parser.add_argument("--docx-title", default="参考文献", help="Title used in the exported .docx.")
    parser.add_argument("--docx-font-en", default="Times New Roman", help="Latin font for .docx.")
    parser.add_argument("--docx-font-cn", default="宋体", help="CJK font for .docx (EastAsia).")
    parser.add_argument("--docx-font-size", type=int, default=12, help="Font size (pt) for .docx.")
    parser.add_argument(
        "--docx-hanging-cm",
        type=float,
        default=0.74,
        help="Hanging indent size (cm) for .docx.",
    )

    args = parser.parse_args(argv)

    md_path = Path(args.md_in)
    decoded = _decode_text_file(md_path, args.encoding)

    lines = decoded.text.splitlines()
    heading_idx, section_start, section_end = _find_reference_section(lines, args.heading_regex)
    raw_section_lines = lines[section_start:section_end]

    entries = _parse_entries(raw_section_lines)
    cn_entries = [e for e in entries if _is_cjk_entry(e)]
    en_entries = sorted([e for e in entries if not _is_cjk_entry(e)], key=_sort_key_english)
    formatted_entries = cn_entries + en_entries

    new_section_lines = _format_reference_section(formatted_entries)
    new_lines = lines[: heading_idx + 1] + new_section_lines + lines[section_end:]
    new_text = decoded.newline.join(new_lines).rstrip() + decoded.newline

    if args.write_md:
        md_path.write_text(new_text, encoding=decoded.encoding)
    elif args.md_out:
        Path(args.md_out).write_text(new_text, encoding=decoded.encoding)

    if args.docx_out:
        _export_docx(
            entries=formatted_entries,
            out_path=Path(args.docx_out),
            title=args.docx_title,
            font_en=args.docx_font_en,
            font_cn=args.docx_font_cn,
            font_size_pt=args.docx_font_size,
            hanging_cm=args.docx_hanging_cm,
        )

    print(
        f"Reference entries: total={len(entries)} cn={len(cn_entries)} en={len(en_entries)} "
        f"(written_md={bool(args.write_md or args.md_out)} docx={bool(args.docx_out)})"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
