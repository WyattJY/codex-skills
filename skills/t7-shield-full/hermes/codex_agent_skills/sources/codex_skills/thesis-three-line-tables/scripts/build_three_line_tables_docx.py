#!/usr/bin/env python3
"""
Build a Word (.docx) document containing thesis three-line tables (三线表) from per-table CSVs.

Features:
- No vertical lines, no internal grid lines
- Thick top/bottom rules, thin header separator rule
- Header bold + centered
- Numeric columns right-aligned (heuristic)
- Optional: extract table captions from a Markdown thesis draft (lines like: "表5-2 ...")
- Built-in beautification for long-text API endpoint tables
"""

from __future__ import annotations

import argparse
import csv
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple


try:
    from docx import Document
    from docx.enum.section import WD_SECTION, WD_ORIENTATION
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency: python-docx. Install with: python -m pip install python-docx") from exc


TABLE_CAPTION_RE = re.compile(r"^表(?P<chapter>\d+)-(?P<index>\d+)\s+(?P<title>.+?)\s*$")


@dataclass(frozen=True)
class TableItem:
    caption: str
    table_id: Optional[Tuple[int, int]]
    csv_path: Path


def is_number(value: str) -> bool:
    try:
        float(value)
        return True
    except Exception:
        return False


def read_csv(path: Path) -> List[List[str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.reader(f))


def find_csv_for_table_id(csv_dir: Path, chapter: int, index: int) -> Path:
    patterns = [
        f"table{chapter}_{index}_*.csv",
        f"table{chapter}_{index}.csv",
    ]
    matches: list[Path] = []
    for pattern in patterns:
        matches.extend(csv_dir.glob(pattern))

    matches = sorted(set(matches))
    if len(matches) == 1:
        return matches[0]
    if not matches:
        raise FileNotFoundError(f"No CSV found for 表{chapter}-{index} in {csv_dir} (expected {patterns}).")
    raise FileExistsError(f"Multiple CSVs found for 表{chapter}-{index}: {[m.name for m in matches]}")


def parse_tables_from_markdown(md_path: Path, csv_dir: Path) -> List[TableItem]:
    items: list[TableItem] = []
    for line in md_path.read_text(encoding="utf-8").splitlines():
        m = TABLE_CAPTION_RE.match(line.strip())
        if not m:
            continue
        chapter = int(m.group("chapter"))
        index = int(m.group("index"))
        csv_path = find_csv_for_table_id(csv_dir, chapter=chapter, index=index)
        items.append(TableItem(caption=line.strip(), table_id=(chapter, index), csv_path=csv_path))

    if not items:
        raise ValueError(f"No table captions found in Markdown: {md_path}")
    return items


def table_is_api_endpoints(item: TableItem) -> bool:
    if item.table_id == (5, 2):
        return True
    name = item.csv_path.name.lower()
    if "api" in name or "endpoint" in name:
        return True
    if "API" in item.caption or "端点" in item.caption:
        return True
    return False


def set_run_font(run, size_pt: float, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph(paragraph, text: str, size_pt: float, bold: bool, align) -> None:
    paragraph.text = ""
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold)


def set_table_borders_none(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def set_cell_border(cell, *, top="nil", bottom="nil", top_sz=0, bottom_sz=0, color="000000") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    def set_edge(edge: str, val: str, sz: int) -> None:
        existing = tc_borders.find(qn("w:" + edge))
        if existing is not None:
            tc_borders.remove(existing)
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), val)
        if val != "nil":
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        tc_borders.append(el)

    set_edge("top", top, top_sz if top != "nil" else 0)
    set_edge("bottom", bottom, bottom_sz if bottom != "nil" else 0)
    set_edge("left", "nil", 0)
    set_edge("right", "nil", 0)


def apply_three_line_style(table, thick_pt: float, mid_pt: float) -> None:
    thick_sz = int(round(thick_pt * 8))  # Word border sz is 1/8 pt
    mid_sz = int(round(mid_pt * 8))

    set_table_borders_none(table)
    n_rows = len(table.rows)
    n_cols = len(table.columns)

    for r in range(n_rows):
        for c in range(n_cols):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if r == 0:
                set_cell_border(cell, top="single", bottom="single", top_sz=thick_sz, bottom_sz=mid_sz)
            elif r == n_rows - 1:
                set_cell_border(cell, top="nil", bottom="single", bottom_sz=thick_sz)
            else:
                set_cell_border(cell, top="nil", bottom="nil")


def guess_column_alignment(header: List[str], data: List[List[str]]) -> List[int]:
    aligns: list[int] = []
    for col_idx, _ in enumerate(header):
        values = [r[col_idx] for r in data if col_idx < len(r) and r[col_idx] != ""]
        if not values:
            aligns.append(WD_ALIGN_PARAGRAPH.LEFT)
            continue
        numeric_ratio = sum(is_number(v) for v in values) / len(values)
        aligns.append(WD_ALIGN_PARAGRAPH.RIGHT if numeric_ratio >= 0.8 else WD_ALIGN_PARAGRAPH.LEFT)
    return aligns


def format_api_table_cell(col_idx: int, value: str) -> str:
    if col_idx == 2:
        return value.replace("；", "\n").replace(";", "\n")
    if col_idx == 3 and "（" in value and len(value) > 18:
        left, right = value.split("（", 1)
        return left.strip() + "\n（" + right
    return value


def build_docx(
    items: List[TableItem],
    output: Path,
    title: str,
    landscape_threshold: int,
    normal_font_size: float,
    wide_font_size: float,
    thick_pt: float,
    mid_pt: float,
) -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(normal_font_size)

    title_p = doc.add_paragraph()
    set_paragraph(title_p, title, size_pt=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    portrait_w = doc.sections[0].page_width
    portrait_h = doc.sections[0].page_height
    landscape_w = portrait_h
    landscape_h = portrait_w

    for item in items:
        rows = read_csv(item.csv_path)
        if not rows:
            raise ValueError(f"Empty CSV: {item.csv_path}")
        header = rows[0]
        data = rows[1:]

        is_api = table_is_api_endpoints(item)
        is_wide = len(header) >= landscape_threshold
        landscape = is_api or is_wide

        section = doc.add_section(WD_SECTION.NEW_PAGE)
        if landscape:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = landscape_w
            section.page_height = landscape_h
        else:
            section.orientation = WD_ORIENTATION.PORTRAIT
            section.page_width = portrait_w
            section.page_height = portrait_h

        cap_p = doc.add_paragraph()
        set_paragraph(cap_p, item.caption, size_pt=normal_font_size, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

        table = doc.add_table(rows=len(data) + 1, cols=len(header))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        apply_three_line_style(table, thick_pt=thick_pt, mid_pt=mid_pt)

        font_size = wide_font_size if landscape else normal_font_size

        aligns = guess_column_alignment(header, data)
        if is_api and len(header) == 4:
            # prettier layout
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
            widths = [Cm(2.4), Cm(2.0), Cm(9.0), Cm(12.0)]
            for col_idx, w in enumerate(widths):
                table.columns[col_idx].width = w
                for r in table.rows:
                    r.cells[col_idx].width = w

        # Header
        for j, col_name in enumerate(header):
            set_paragraph(table.cell(0, j).paragraphs[0], col_name, size_pt=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Data
        for i, row in enumerate(data, start=1):
            for j, value in enumerate(row):
                if is_api:
                    value = format_api_table_cell(j, value)
                set_paragraph(table.cell(i, j).paragraphs[0], value, size_pt=font_size, bold=False, align=aligns[j])

        # Optional note for ablation tables
        if item.table_id == (3, 2) or "消融" in item.caption:
            note_p = doc.add_paragraph()
            set_paragraph(note_p, "注：× 表示未使用该模块，√ 表示使用该模块。", size_pt=wide_font_size, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", type=Path, required=True, help="Folder containing per-table CSV files.")
    parser.add_argument("--output", type=Path, required=True, help="Output .docx path.")
    parser.add_argument("--md", type=Path, default=None, help="Optional Markdown file to extract captions like '表5-2 ...'.")
    parser.add_argument("--title", default="三线表汇总（可直接复制到论文 Word）", help="Document title (first page).")
    parser.add_argument("--landscape-threshold", type=int, default=8, help="Landscape if column count >= this value.")
    parser.add_argument("--font-size", type=float, default=10.5, help="Normal table font size (pt).")
    parser.add_argument("--wide-font-size", type=float, default=9.0, help="Wide/landscape table font size (pt).")
    parser.add_argument("--thick-pt", type=float, default=2.25, help="Top/bottom rule thickness (pt).")
    parser.add_argument("--mid-pt", type=float, default=1.0, help="Header separator thickness (pt).")
    args = parser.parse_args()

    if args.md is not None:
        items = parse_tables_from_markdown(args.md, csv_dir=args.csv_dir)
    else:
        csvs = sorted(args.csv_dir.glob("*.csv"))
        if not csvs:
            raise SystemExit(f"No CSV files found in: {args.csv_dir}")
        items = [TableItem(caption=p.stem, table_id=None, csv_path=p) for p in csvs]

    build_docx(
        items=items,
        output=args.output,
        title=args.title,
        landscape_threshold=args.landscape_threshold,
        normal_font_size=args.font_size,
        wide_font_size=args.wide_font_size,
        thick_pt=args.thick_pt,
        mid_pt=args.mid_pt,
    )
    print(str(args.output))


if __name__ == "__main__":
    main()

